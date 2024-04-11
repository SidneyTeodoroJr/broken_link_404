import requests
from bs4 import BeautifulSoup
from docx import Document

# Função para verificar links quebrados em uma página HTML
def verificar_links_quebrados(url):
    try:
        response = requests.get(url)
        response.raise_for_status()  # Levanta uma exceção se a requisição falhar
        soup = BeautifulSoup(response.text, 'html.parser')
        links_quebrados = []

        # Encontrar todos os elementos 'a' (links) na página
        for link in soup.find_all('a', href=True):
            href = link['href']
            if not href.startswith(('http://', 'https://')):
                url_completa = url.rstrip('/') + '/' + href.lstrip('/')
                if requests.head(url_completa).status_code == 404:
                    links_quebrados.append((href, url_completa, str(link)))

        return links_quebrados

    except Exception as e:
        print(f"Erro ao acessar {url}: {e}")
        return []

# Função para salvar os links quebrados em um documento do Word
def salvar_links_quebrados(links_quebrados, nome_arquivo):
    doc = Document()
    doc.add_heading('''Links Quebrados:
                    ''', level=1)

    for link in links_quebrados:
        doc.add_paragraph(f'Link: {link[0]}')
        doc.add_paragraph(f'URL: {link[1]}')
        doc.add_paragraph(f'Elemento HTML: {link[2]}')
        doc.add_paragraph('---')

    doc.save(nome_arquivo)

# URL da página que deseja verificar os links
url_pagina = ''

# Verificar os links quebrados
links_quebrados = verificar_links_quebrados(url_pagina)

# Salvar os links quebrados em um documento do Word
nome_arquivo = 'app\doc\links_quebrados.docx'
salvar_links_quebrados(links_quebrados, nome_arquivo)

print(f'Os links quebrados foram salvos em {nome_arquivo}')